from __future__ import annotations

"""semiconductor_doc_ingest.py

目标：支持“非结构化文档 -> 三元组(实体/关系/属性) -> 人工确认 -> 增量写入Neo4j”的自动化管道。

你会得到什么：
1) 读取 PDF / DOCX / TXT 等文档，抽取正文文本（DOCX包含表格）
2) 调用大模型（DeepSeek）自动抽取：实体、关系（三元组）、量化知识（工艺窗口/曲线/范围）
3) 输出“待人工确认”的 JSON 文件；你确认后再执行导入
4) 增量更新与版本管理：每次导入会记录 source_doc、source_version、ingest_time

依赖：
- pip install openai neo4j pypdf python-docx

用法：
1) 抽取（不入库）
   python semiconductor_doc_ingest.py --file "D:\\Desktop\\半导体沉积工艺.docx" --source_version "v1" --out triples_pending.json

2) 人工确认后入库
   python semiconductor_doc_ingest.py --apply --pending triples_pending.json

"""

import argparse
import json
import os
import re
import time
from dataclasses import dataclass
from datetime import datetime, timezone
from typing import Any, Dict, List, Tuple
import sys

from neo4j import GraphDatabase, Driver
from openai import OpenAI

import settings

settings.validate_llm_config()
settings.validate_neo4j_config()

NEO4J_URI = settings.NEO4J_URI
NEO4J_USER = settings.NEO4J_USER
NEO4J_PASSWORD = settings.NEO4J_PASSWORD
DEEPSEEK_API_KEY = settings.DEEPSEEK_API_KEY
DEEPSEEK_API_BASE = settings.DEEPSEEK_API_BASE

client = OpenAI(api_key=DEEPSEEK_API_KEY, base_url=DEEPSEEK_API_BASE)

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
DEFAULT_TECH_DATA_PATH = os.path.join(BASE_DIR, "tech_Data.py")

# ---------------------------------------------------------------------------
# 图谱标签（与 tech_Data.py 对齐）
# ---------------------------------------------------------------------------
LABELS = [
    "Action",
    "Technology",
    "Method",
    "SubMethod",
    "Material",
    "Capability",
    "Equipment",
    "ChipStructure",
    "ManufacturingStage",
    "Parameter",
]


@dataclass
class Entity:
    label: str
    name: str
    description: str = ""
    properties: Dict[str, Any] | None = None


@dataclass
class Relation:
    start_label: str
    start_name: str
    rel_type: str
    end_label: str
    end_name: str
    description: str = ""
    properties: Dict[str, Any] | None = None


@dataclass
class ExtractionPackage:
    source_doc: str
    source_version: str
    extracted_at: str
    entities: List[Entity]
    relations: List[Relation]
    measurements: List[Dict[str, Any]]
    warnings: List[str]


# ---------------------------------------------------------------------------
# 文档解析：PDF / DOCX / TXT
# ---------------------------------------------------------------------------

def read_pdf_text(file_path: str, max_pages: int = 50) -> str:
    try:
        from pypdf import PdfReader
    except Exception as e:  # pragma: no cover
        raise RuntimeError("请先安装 pypdf：pip install pypdf") from e

    reader = PdfReader(file_path)
    texts: List[str] = []
    pages_to_read = min(len(reader.pages), max_pages)
    for i in range(pages_to_read):
        t = (reader.pages[i].extract_text() or "").strip()
        if t:
            texts.append(t)
    return "\n\n".join(texts)


def read_docx_text(file_path: str) -> str:
    try:
        import docx
    except Exception as e:  # pragma: no cover
        raise RuntimeError("请先安装 python-docx：pip install python-docx") from e

    doc = docx.Document(file_path)
    parts: List[str] = []

    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if t:
            parts.append(t)

    # 表格：每行拼成一行文本（便于抽取）
    for table in doc.tables:
        for row in table.rows:
            cells = [(c.text or "").strip() for c in row.cells]
            cells = [c for c in cells if c]
            if cells:
                parts.append(" | ".join(cells))

    return "\n".join(parts)


def read_text_file(file_path: str) -> str:
    with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
        return f.read()


def load_document(file_path: str) -> str:
    ext = os.path.splitext(file_path)[1].lower()
    if ext == ".pdf":
        return read_pdf_text(file_path)
    if ext == ".docx":
        return read_docx_text(file_path)
    return read_text_file(file_path)


# ---------------------------------------------------------------------------
# JSON 解析健壮性：提取 + 多次修复
# ---------------------------------------------------------------------------

def _extract_json_block(text: str) -> str:
    """从模型输出中尽量提取 JSON 对象文本。"""
    text = text.strip()
    # 查找被 ```json ... ``` 包裹的代码块
    match = re.search(r"```json\n(\{[\s\S]*\})\n```", text)
    if match:
        return match.group(1)
    
    # 查找第一个 '{' 和最后一个 '}' 之间的内容
    match = re.search(r"\{[\s\S]*\}", text)
    if match:
        return match.group(0)

    raise ValueError("模型输出中未找到有效的JSON对象")


def _llm_fix_json(bad_json_text: str, attempt: int) -> Dict[str, Any]:
    """让模型把‘不合法/夹杂文字的输出’修复为严格 JSON。"""
    if attempt == 1:
        instruction = "请你只做一件事：把它修复成【严格合法的 JSON】并输出。如果原文被截断（比如最后一个大括号缺失），请合理地补全，确保语法正确。"
    else:
        instruction = "你只需要从下面的文本里提取出JSON部分，别的什么都不要做。"

    fix_prompt = f"""你是JSON修复器。

下面是一段内容，其中可能包含JSON，也可能是接近JSON但不合法。
{instruction}

要求：
- 只能输出 JSON，不要任何解释、不要 Markdown。
- 不要新增任何原文中不存在的新事实/新实体/新关系（禁止编造）。
- 如果字段缺失，补空数组：entities/relations/measurements/warnings。

待修复内容：
---
{bad_json_text}
---
"""

    resp = client.chat.completions.create(
        model="deepseek-chat",
        messages=[{"role": "user", "content": fix_prompt}],
        temperature=0.0,
        max_tokens=4000,
    )

    content = (resp.choices[0].message.content or "").strip()
    json_str = _extract_json_block(content)
    return json.loads(json_str)


# ---------------------------------------------------------------------------
# LLM 抽取（反幻觉 + 强制 evidence）
# ---------------------------------------------------------------------------

def _llm_extract_triples(text: str) -> Dict[str, Any]:
    """调用 LLM 抽取结构化结果，并具备多次修复能力。"""
    text = re.sub(r"\s+", " ", text)[:12000]

    prompt = f"""你是“半导体知识图谱构建器”。

从下列文本中抽取实体、关系（三元组）、量化知识。

【强约束（用来避免幻觉）】
1) label 只能从以下集合中选择：{', '.join(LABELS)}
2) 每一条 entity / relation / measurement 都必须包含 evidence（原文可逐字找到的一段短文本）。
   - 如果找不到 evidence：不要输出该条。
3) 宁可少抽取，也不要编造设备型号、参数范围、公式、单位。
4) 输出必须是【严格合法 JSON】，用 ```json ... ``` 包裹。

【输出 JSON 结构】
```json
{{
  "entities": [
    {{"label":"Technology","name":"PVD","description":"...","properties":{{}},"evidence":"..."}}
  ],
  "relations": [
    {{"start_label":"Method","start_name":"溅射","rel_type":"NEED_EQUIPMENT","end_label":"Equipment","end_name":"溅射台","description":"...","properties":{{}},"evidence":"..."}}
  ],
  "measurements": [
    {{"subject_label":"Parameter","subject_name":"温度","metric":"沉积速率","conditions":{{}},"value":{{}},"evidence":"..."}}
  ],
  "warnings": []
}}
```

【待抽取文本】
{text}
"""

    resp = client.chat.completions.create(
        model="deepseek-chat",
        messages=[{"role": "user", "content": prompt}],
        temperature=0.1,
        max_tokens=4000,
    )

    content = (resp.choices[0].message.content or "").strip()

    try:
        with open("llm_raw_output.txt", "w", encoding="utf-8") as f:
            f.write(content)
    except Exception:
        pass

    # 两阶段修复，最大化健壮性
    try:
        json_str = _extract_json_block(content)
        return json.loads(json_str)
    except Exception as e1:
        print(f"第一次JSON解析失败: {e1}。尝试让LLM修复...")
        try:
            return _llm_fix_json(content, 1)
        except Exception as e2:
            print(f"第二次JSON修复失败: {e2}。尝试终极提取...")
            try:
                return _llm_fix_json(content, 2)
            except Exception as e3:
                print(f"最终JSON提取失败: {e3}。本次抽取返回空结果。")
                with open("llm_bad_output.txt", "w", encoding="utf-8") as f:
                    f.write(content)
                return {"entities": [], "relations": [], "measurements": [], "warnings": ["LLM返回内容无法解析为JSON"]}


# ---------------------------------------------------------------------------
# 组装 ExtractionPackage（并做 evidence 强制过滤）
# ---------------------------------------------------------------------------

def _dedup_entities(entities: List[Entity]) -> List[Entity]:
    seen: set[Tuple[str, str]] = set()
    out: List[Entity] = []
    for e in entities:
        key = (e.label, e.name.strip())
        if not e.label or not e.name or e.label not in LABELS or key in seen:
            continue
        seen.add(key)
        out.append(e)
    return out


def _dedup_relations(relations: List[Relation]) -> List[Relation]:
    seen: set[Tuple[str, str, str, str, str]] = set()
    out: List[Relation] = []
    for r in relations:
        key = (r.start_label, r.start_name, r.rel_type, r.end_label, r.end_name)
        if not all(key) or r.start_label not in LABELS or r.end_label not in LABELS or key in seen:
            continue
        seen.add(key)
        out.append(r)
    return out


def extract_from_document(file_path: str, source_version: str) -> ExtractionPackage:
    text = load_document(file_path)
    if not text.strip():
        raise RuntimeError("文档解析为空：可能是扫描PDF或DOCX内容为空")

    data = _llm_extract_triples(text)

    entities: List[Entity] = []
    for e in data.get("entities", []) or []:
        evidence = (e.get("evidence") or "").strip()
        if not evidence:
            continue
        entities.append(
            Entity(
                label=e.get("label", ""),
                name=e.get("name", ""),
                description=e.get("description", ""),
                properties={**(e.get("properties") or {}), "evidence": evidence},
            )
        )

    relations: List[Relation] = []
    for r in data.get("relations", []) or []:
        evidence = (r.get("evidence") or "").strip()
        if not evidence:
            continue
        relations.append(
            Relation(
                start_label=r.get("start_label", ""),
                start_name=r.get("start_name", ""),
                rel_type=r.get("rel_type", ""),
                end_label=r.get("end_label", ""),
                end_name=r.get("end_name", ""),
                description=r.get("description", ""),
                properties={**(r.get("properties") or {}), "evidence": evidence},
            )
        )

    measurements = [
        m
        for m in (data.get("measurements", []) or [])
        if (m.get("evidence") or "").strip()
    ]

    return ExtractionPackage(
        source_doc=os.path.basename(file_path),
        source_version=source_version,
        extracted_at=datetime.now(timezone.utc).isoformat().replace("+00:00", "Z"),
        entities=_dedup_entities(entities),
        relations=_dedup_relations(relations),
        measurements=measurements,
        warnings=data.get("warnings", [])
    )


# ---------------------------------------------------------------------------
# 人工确认输出（JSON）
# ---------------------------------------------------------------------------

def save_pending(pkg: ExtractionPackage, out_path: str) -> None:
    payload = {
        "source_doc": pkg.source_doc,
        "source_version": pkg.source_version,
        "extracted_at": pkg.extracted_at,
        "entities": [
            {
                "label": e.label,
                "name": e.name,
                "description": e.description,
                "properties": e.properties or {},
                "confirmed": False,
            }
            for e in pkg.entities
        ],
        "relations": [
            {
                "start_label": r.start_label,
                "start_name": r.start_name,
                "rel_type": r.rel_type,
                "end_label": r.end_label,
                "end_name": r.end_name,
                "description": r.description,
                "properties": r.properties or {},
                "confirmed": False,
            }
            for r in pkg.relations
        ],
        "measurements": [{**m, "confirmed": False} for m in (pkg.measurements or [])],
        "warnings": pkg.warnings,
    }

    with open(out_path, "w", encoding="utf-8") as f:
        json.dump(payload, f, ensure_ascii=False, indent=2)


def load_pending(pending_path: str) -> Dict[str, Any]:
    with open(pending_path, "r", encoding="utf-8") as f:
        return json.load(f)


# ---------------------------------------------------------------------------
# 同步写入 tech_Data.py（让每次提取都落盘到技术数据）
# ---------------------------------------------------------------------------

LABEL_TO_LIST = {
    "Action": "Action",
    "Technology": "Technology",
    "Method": "Method",
    "SubMethod": "SubMethod",
    "Material": "Material",
    "Capability": "Capability",
    "Equipment": "Equipment",
    "ChipStructure": "ChipStructure",
    "ManufacturingStage": "ManufacturingStage",
    "Parameter": "Parameter",
}

ID_PREFIX = {
    "Action": "action",
    "Technology": "tech",
    "Method": "method",
    "SubMethod": "submethod",
    "Material": "material",
    "Capability": "capability",
    "Equipment": "equipment",
    "ChipStructure": "structure",
    "ManufacturingStage": "stage",
    "Parameter": "parameter",
}


def _find_list_block(text: str, anchor: str) -> Tuple[int, int, str]:
    start_anchor = text.find(anchor)
    if start_anchor < 0:
        raise ValueError(f"在 tech_Data.py 中未找到锚点: {anchor}")

    start_bracket = text.find("[", start_anchor)
    if start_bracket < 0:
        raise ValueError(f"在 tech_Data.py 中未找到列表起始 '[': {anchor}")

    depth = 0
    i = start_bracket
    while i < len(text):
        ch = text[i]
        if ch == "[":
            depth += 1
        elif ch == "]":
            depth -= 1
            if depth == 0:
                return start_bracket, i, text[start_bracket:i + 1]
        i += 1

    raise ValueError(f"在 tech_Data.py 中未找到列表结束 ']': {anchor}")


def _next_id(existing_ids: List[str], prefix: str) -> str:
    pat = re.compile(rf"^{re.escape(prefix)}_(\d+)$")
    nums = []
    for eid in existing_ids:
        m = pat.match(eid)
        if m:
            nums.append(int(m.group(1)))
    nxt = (max(nums) + 1) if nums else 1
    return f"{prefix}_{nxt:03d}"


def _sync_extraction_to_tech_data(pkg: ExtractionPackage, tech_data_path: str = "tech_Data.py") -> Dict[str, int]:
    with open(tech_data_path, "r", encoding="utf-8") as f:
        text = f.read()

    # 1) 读取现有实体，构建 name->id 索引
    entity_index: Dict[str, Dict[str, str]] = {k: {} for k in LABEL_TO_LIST}
    entity_ids: Dict[str, List[str]] = {k: [] for k in LABEL_TO_LIST}

    for label, list_name in LABEL_TO_LIST.items():
        _, _, block = _find_list_block(text, f"self.{list_name} = [")
        for m in re.finditer(r'\{"id":\s*"([^"]+)",\s*"name":\s*"([^"]+)"', block):
            eid, name = m.group(1), m.group(2)
            entity_index[label][name] = eid
            entity_ids[label].append(eid)

    added_entities = 0
    skipped_entities = 0

    # 2) 增量追加实体
    for label, list_name in LABEL_TO_LIST.items():
        extracted = [e for e in pkg.entities if e.label == label and e.name.strip()]
        if not extracted:
            continue

        start, end, _ = _find_list_block(text, f"self.{list_name} = [")
        insert_lines: List[str] = []

        for e in extracted:
            name = e.name.strip()
            if name in entity_index[label]:
                skipped_entities += 1
                continue

            new_id = _next_id(entity_ids[label], ID_PREFIX[label])
            desc = (e.description or "").strip().replace("\n", " ")
            line = f'            {{"id": "{new_id}", "name": {json.dumps(name, ensure_ascii=False)}, "description": {json.dumps(desc, ensure_ascii=False)}}},'
            insert_lines.append(line)
            entity_index[label][name] = new_id
            entity_ids[label].append(new_id)
            added_entities += 1

        if insert_lines:
            insertion = "\n" + "\n".join(insert_lines)
            text = text[:end] + insertion + text[end:]

    # 3) 增量追加关系
    r_start, r_end, r_block = _find_list_block(text, "self.relationships = [")

    existing_rel_keys = set()
    for m in re.finditer(
        r'\{"start_id":\s*"([^"]+)",\s*"end_id":\s*"([^"]+)",\s*"relationship_type":\s*"([^"]+)"',
        r_block,
    ):
        existing_rel_keys.add((m.group(1), m.group(2), m.group(3)))

    added_relations = 0
    skipped_relations = 0
    relation_lines: List[str] = []

    for r in pkg.relations:
        s_name = r.start_name.strip()
        e_name = r.end_name.strip()
        s_label = r.start_label.strip()
        e_label = r.end_label.strip()

        if s_label not in entity_index or e_label not in entity_index:
            skipped_relations += 1
            continue

        s_id = entity_index[s_label].get(s_name)
        e_id = entity_index[e_label].get(e_name)
        if not s_id or not e_id:
            skipped_relations += 1
            continue

        rel_type = (r.rel_type or "RELATED_TO").strip() or "RELATED_TO"
        key = (s_id, e_id, rel_type)
        if key in existing_rel_keys:
            skipped_relations += 1
            continue

        desc = (r.description or "").strip().replace("\n", " ")
        line = (
            f'            {{"start_id": "{s_id}", "end_id": "{e_id}", '
            f'"relationship_type": "{rel_type}", "weight": 1, '
            f'"description": {json.dumps(desc, ensure_ascii=False)}}},'
        )
        relation_lines.append(line)
        existing_rel_keys.add(key)
        added_relations += 1

    if relation_lines:
        text = text[:r_end] + "\n" + "\n".join(relation_lines) + text[r_end:]

    with open(tech_data_path, "w", encoding="utf-8") as f:
        f.write(text)

    return {
        "added_entities": added_entities,
        "skipped_entities": skipped_entities,
        "added_relations": added_relations,
        "skipped_relations": skipped_relations,
    }


# ---------------------------------------------------------------------------
# 增量导入 Neo4j（仅导入 confirmed=true 的条目）
# ---------------------------------------------------------------------------

def _merge_entity(tx, e: Dict[str, Any], meta: Dict[str, Any]) -> None:
    label = e["label"]
    name = e["name"]
    desc = e.get("description") or ""
    props = e.get("properties") or {}

    props_meta = {
        **props,
        "source_doc": meta["source_doc"],
        "source_version": meta["source_version"],
        "ingest_time": meta["ingest_time"],
    }

    query = f"""
    MERGE (n:{label} {{name: $name}})
    ON CREATE SET n.id = coalesce(n.id, $id), n.description = $desc
    SET n += $props
    """

    tx.run(query, name=name, id=e.get("id"), desc=desc, props=props_meta)


def _merge_relation(tx, r: Dict[str, Any], meta: Dict[str, Any]) -> None:
    start_label = r["start_label"]
    start_name = r["start_name"]
    end_label = r["end_label"]
    end_name = r["end_name"]
    rel_type = r["rel_type"]

    desc = r.get("description") or ""
    props = r.get("properties") or {}

    props_meta = {
        **props,
        "description": desc,
        "source_doc": meta["source_doc"],
        "source_version": meta["source_version"],
        "ingest_time": meta["ingest_time"],
    }

    query = f"""
    MATCH (a:{start_label} {{name:$start_name}})
    MATCH (b:{end_label} {{name:$end_name}})
    MERGE (a)-[r:{rel_type}]->(b)
    SET r += $props
    """

    tx.run(query, start_name=start_name, end_name=end_name, props=props_meta)


def apply_pending_to_neo4j(pending: Dict[str, Any]) -> None:
    meta = {
        "source_doc": pending.get("source_doc", "unknown"),
        "source_version": pending.get("source_version", "unknown"),
        "ingest_time": datetime.now(timezone.utc).isoformat().replace("+00:00", "Z"),
    }

    driver: Driver = GraphDatabase.driver(NEO4J_URI, auth=(NEO4J_USER, NEO4J_PASSWORD))
    try:
        with driver.session() as session:
            for e in pending.get("entities", []):
                if not e.get("confirmed"):
                    continue
                session.execute_write(_merge_entity, e, meta)

            for r in pending.get("relations", []):
                if not r.get("confirmed"):
                    continue
                session.execute_write(_merge_relation, r, meta)

        print("已导入已确认的实体与关系到 Neo4j。")
    finally:
        driver.close()


# ---------------------------------------------------------------------------
# CLI
# ---------------------------------------------------------------------------

def main() -> None:
    parser = argparse.ArgumentParser(description="半导体文档导入：PDF/DOCX/TXT -> 三元组 -> 人工确认 -> 增量入库")

    parser.add_argument("--file", type=str, help="要解析的文档路径（pdf/docx/txt等）")
    parser.add_argument("--source_version", type=str, default="v1", help="文档版本号（用于溯源）")
    parser.add_argument("--out", type=str, default="triples_pending.json", help="输出待确认JSON")
    parser.add_argument("--tech_data", type=str, default=DEFAULT_TECH_DATA_PATH, help="tech_Data.py 文件路径（每次提取都会增量写入）")

    parser.add_argument("--apply", action="store_true", help="将 pending.json 中 confirmed=true 的内容写入 Neo4j")
    parser.add_argument("--pending", type=str, default="triples_pending.json", help="待确认JSON文件路径")

    args = parser.parse_args()

    if args.apply:
        pending = load_pending(args.pending)
        apply_pending_to_neo4j(pending)
        return

    if not args.file:
        raise SystemExit("请提供 --file 文档路径；或使用 --apply --pending 进行导入")

    print("开始解析文档并抽取三元组...")
    start = time.time()

    try:
        pkg = extract_from_document(args.file, args.source_version)
        save_pending(pkg, args.out)

        sync_stat = _sync_extraction_to_tech_data(pkg, args.tech_data)

        print(f"完成：抽取实体 {len(pkg.entities)} 个，关系 {len(pkg.relations)} 条，量化信息 {len(pkg.measurements)} 条")
        print(f"已输出待人工确认文件：{args.out}")
        print(
            "已同步写入 tech_Data："
            f"新增实体 {sync_stat['added_entities']}，"
            f"跳过实体 {sync_stat['skipped_entities']}，"
            f"新增关系 {sync_stat['added_relations']}，"
            f"跳过关系 {sync_stat['skipped_relations']}"
        )
    except Exception as e:
        print(f"文档抽取失败: {e}", file=sys.stderr)

    print(f"耗时：{time.time() - start:.2f}s")


if __name__ == "__main__":
    main()
